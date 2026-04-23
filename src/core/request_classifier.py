"""Unified request classifier — Phase 1 of the PC↔RFQ refactor.

Takes any buyer request (email + attachments, manual upload, or API)
and returns a `RequestClassification` describing:

  shape    — what kind of document structure this is (CCHCS packet,
             DOCX 704, Docusign flat 704, XLSX RFQ, etc.)
  agency   — which CA state agency the request is for
  required_forms — the forms that must be generated in response
                   (sourced from agency_config for the matched agency)
  confidence — 0.0-1.0 — how sure we are
  reasons  — human-readable explanations of every signal that fired

This is the single source of truth for "what is this request and how
should we respond to it?" Every ingest path (email poller, manual
upload, /api/v1) routes through `classify_request()`. The result
drives downstream dispatch in `ingest_pipeline.py` (Phase 2) and the
detail-page conditional UI (Phase 3).

Feature-flagged via `ingest.classifier_v2_enabled` — callers can
disable it to fall back to the legacy parallel PC/RFQ paths if
anything goes sideways.

Ground truth: tests in test_request_classifier.py exercise every
classification branch against real fixtures in
`tests/fixtures/unified_ingest/`.
"""
from __future__ import annotations

import logging
import os
import re
from dataclasses import dataclass, field
from typing import Any, Dict, List, Optional, Tuple

log = logging.getLogger("reytech.classifier")


# ── Shape taxonomy ────────────────────────────────────────────────────────

SHAPE_CCHCS_PACKET = "cchcs_packet"            # 18-page CCHCS buyer-issued packet
SHAPE_CCHCS_IT_RFQ = "cchcs_it_rfq"            # CCHCS LPA IT Goods & Services RFQ (Non-Cloud)
SHAPE_PC_704_DOCX = "pc_704_docx"              # DOCX AMS 704 worksheet
SHAPE_PC_704_PDF_DOCUSIGN = "pc_704_pdf_docusign"  # flat DocuSign PDF (overlay path)
SHAPE_PC_704_PDF_FILLABLE = "pc_704_pdf_fillable"  # fillable PDF (form-field path)
SHAPE_GENERIC_RFQ_XLSX = "generic_rfq_xlsx"    # spreadsheet RFQ
SHAPE_GENERIC_RFQ_PDF = "generic_rfq_pdf"      # non-704 PDF RFQ
SHAPE_GENERIC_RFQ_DOCX = "generic_rfq_docx"    # non-704 DOCX RFQ
SHAPE_EMAIL_ONLY = "email_only"                # no attachments — body only
SHAPE_UNKNOWN = "unknown"                      # cannot classify

ALL_SHAPES = {
    SHAPE_CCHCS_PACKET,
    SHAPE_CCHCS_IT_RFQ,
    SHAPE_PC_704_DOCX,
    SHAPE_PC_704_PDF_DOCUSIGN,
    SHAPE_PC_704_PDF_FILLABLE,
    SHAPE_GENERIC_RFQ_XLSX,
    SHAPE_GENERIC_RFQ_PDF,
    SHAPE_GENERIC_RFQ_DOCX,
    SHAPE_EMAIL_ONLY,
    SHAPE_UNKNOWN,
}


# Buyer-supplied templates: the operator can only fill these if the BUYER
# sent one. Agency config may list them as required, but they only actually
# apply when the classifier shape indicates the buyer included them. An LPA
# IT Goods RFQ (shape=email_only / generic_rfq_pdf / generic_rfq_docx) from
# CCHCS must NOT require 703B/704B/bidpkg just because the agency is cchcs —
# the buyer never sent those templates, so Reytech has nothing to fill.
BUYER_TEMPLATE_FORMS = frozenset({"703b", "703c", "704b", "bidpkg"})

# Which buyer-template forms each shape legitimately requires. A shape not
# listed here falls into the "unknown/permissive" bucket — don't strip.
_SHAPE_BUYER_TEMPLATE_REQUIREMENTS: Dict[str, frozenset] = {
    SHAPE_CCHCS_PACKET: frozenset({"703b", "703c", "704b", "bidpkg"}),
    SHAPE_PC_704_DOCX: frozenset({"704b"}),
    SHAPE_PC_704_PDF_DOCUSIGN: frozenset({"704b"}),
    SHAPE_PC_704_PDF_FILLABLE: frozenset({"704b"}),
    SHAPE_GENERIC_RFQ_XLSX: frozenset(),
    SHAPE_GENERIC_RFQ_PDF: frozenset(),
    SHAPE_GENERIC_RFQ_DOCX: frozenset(),
    SHAPE_EMAIL_ONLY: frozenset(),
}


def filter_required_forms_by_shape(required_forms, shape, uploaded_templates=None):
    """Narrow agency-level required_forms to what actually applies for this shape.

    Buyer-template forms (703B/703C/704B/bidpkg) survive the filter only when:
      1. the classifier shape indicates the buyer sent that template, OR
      2. the operator has explicitly uploaded a matching template for this record.

    Non-buyer-template forms (quote, sellers_permit, std204, dvbe843, etc.)
    always pass through — they are Reytech-supplied and unrelated to shape.

    Unknown or missing shapes fall back to "uploaded only" for buyer templates,
    so a manually-created record without stored classification still behaves
    correctly: don't require a 704B unless the operator uploaded one.

    Returns a new list preserving input order.
    """
    if not required_forms:
        return []
    uploaded = {str(k).lower() for k in (uploaded_templates or ())}
    shape_allowed = _SHAPE_BUYER_TEMPLATE_REQUIREMENTS.get(shape or "")
    # 703B and 703C are interchangeable — uploading either satisfies both.
    if "703b" in uploaded or "703c" in uploaded:
        uploaded.update({"703b", "703c"})
    out = []
    for f in required_forms:
        fl = str(f).lower()
        if fl not in BUYER_TEMPLATE_FORMS:
            out.append(f)
            continue
        if shape_allowed is None:
            # Unknown shape → only require if operator uploaded this template.
            if fl in uploaded:
                out.append(f)
            continue
        if fl in shape_allowed or fl in uploaded:
            out.append(f)
    return out


# ── Data model ───────────────────────────────────────────────────────────

@dataclass
class RequestClassification:
    """The canonical answer to 'what is this buyer request?'

    Every downstream step reads from this — no reparsing, no
    re-guessing. The classifier runs once at ingest and stores the
    result on the record.
    """
    shape: str = SHAPE_UNKNOWN
    agency: str = "other"
    agency_name: str = ""
    required_forms: List[str] = field(default_factory=list)
    optional_forms: List[str] = field(default_factory=list)
    confidence: float = 0.0
    reasons: List[str] = field(default_factory=list)
    primary_file: str = ""
    primary_file_type: str = ""
    is_quote_only: bool = False  # True for simple PC worksheets; False for full package RFQs
    needs_overlay_fill: bool = False  # True for flat/DocuSign PDFs
    solicitation_number: str = ""
    institution: str = ""
    producer_signature: str = ""  # PDF /Producer metadata
    meta: Dict[str, Any] = field(default_factory=dict)

    def to_dict(self) -> Dict[str, Any]:
        return {
            "shape": self.shape,
            "agency": self.agency,
            "agency_name": self.agency_name,
            "required_forms": list(self.required_forms),
            "optional_forms": list(self.optional_forms),
            "confidence": self.confidence,
            "reasons": list(self.reasons),
            "primary_file": self.primary_file,
            "primary_file_type": self.primary_file_type,
            "is_quote_only": self.is_quote_only,
            "needs_overlay_fill": self.needs_overlay_fill,
            "solicitation_number": self.solicitation_number,
            "institution": self.institution,
            "producer_signature": self.producer_signature,
            "meta": dict(self.meta),
        }


# ── Agency detection patterns ────────────────────────────────────────────

# Keyword → canonical agency_config key. First match wins per text span.
# Order matters: more specific before less specific (CCHCS before CDCR).
AGENCY_KEYWORDS: List[Tuple[str, str]] = [
    # CCHCS (CA Correctional Health Care Services) — CDCR-adjacent
    (r"\bCALIFORNIA\s+CORRECTIONAL\s+HEALTH\s+CARE\s+SERVICES\b", "cchcs"),
    (r"\bCCHCS\b", "cchcs"),
    (r"\bcchcs\.ca\.gov\b", "cchcs"),
    # CalVet — Barstow first (more specific)
    (r"\bBARSTOW\b", "calvet_barstow"),
    (r"\bVETERANS\s+HOME\s+OF\s+CALIFORNIA\b", "calvet"),
    (r"\bVHC[-\s]?(?:WLA|YV|CH|LB|FR|RED|VRV)\b", "calvet"),
    (r"\bCALVET\b", "calvet"),
    (r"\bCAL\s+VET\b", "calvet"),
    (r"\bCDVA\b", "calvet"),
    (r"\bcalvet\.ca\.gov\b", "calvet"),
    # CDCR (non-CCHCS prisons)
    (r"\bCDCR\b", "cchcs"),  # CDCR mail goes through CCHCS pipeline
    (r"\bCA\s+STATE\s+PRISON\b", "cchcs"),
    (r"\bCORRECTIONS?\s+AND\s+REHABILITATION\b", "cchcs"),
    (r"\bcdcr\.ca\.gov\b", "cchcs"),
    # DSH — State Hospitals
    (r"\bDEPARTMENT\s+OF\s+STATE\s+HOSPITALS\b", "dsh"),
    (r"\bSTATE\s+HOSPITAL\b", "dsh"),
    (r"\bATASCADERO\b|\bCOALINGA\b|\bPATTON\b|\bNAPA\s+STATE\b", "dsh"),
    (r"\bdsh\.ca\.gov\b", "dsh"),
    # CalFire
    (r"\bCAL\s*FIRE\b", "calfire"),
    (r"\bFORESTRY\b", "calfire"),
    (r"\bfire\.ca\.gov\b", "calfire"),
    # DGS (fallback for general services)
    (r"\bDEPARTMENT\s+OF\s+GENERAL\s+SERVICES\b", "dgs"),
    (r"\bdgs\.ca\.gov\b", "dgs"),
]


# Institution prefix → agency. Covers the 2-4 letter facility codes
# that show up in DOCX 704 worksheets and RFQ headers.
INSTITUTION_TO_AGENCY: Dict[str, str] = {
    # CCHCS / CDCR institutions
    "CIM": "cchcs", "CMC": "cchcs", "CIW": "cchcs", "CTF": "cchcs",
    "LAC": "cchcs", "SAC": "cchcs", "SQ": "cchcs", "FSP": "cchcs",
    "SATF": "cchcs", "KVSP": "cchcs", "CRC": "cchcs", "CCWF": "cchcs",
    "CHCF": "cchcs", "DVI": "cchcs", "MCSP": "cchcs", "NKSP": "cchcs",
    "PBSP": "cchcs", "RJD": "cchcs", "SCC": "cchcs", "SOL": "cchcs",
    "SVSP": "cchcs", "VSP": "cchcs", "WSP": "cchcs", "CEN": "cchcs",
    "ISP": "cchcs", "ASP": "cchcs", "HDSP": "cchcs", "CAL": "cchcs",
    "PVSP": "cchcs", "CVSP": "cchcs",
    # CalVet
    "VHC-WLA": "calvet", "VHC-YV": "calvet", "VHC-CH": "calvet",
    "VHC-LB": "calvet", "VHC-FR": "calvet", "VHC-RED": "calvet",
    "VHC-VRV": "calvet",
}


# ── Shape detection signatures ───────────────────────────────────────────

# Strong "this is a CCHCS 18-page packet" signals
CCHCS_PACKET_FIELD_NAMES = {
    "Solicitation No", "Institution Name", "Supplier Name",
    "Item Description1", "Price Per Unit1", "Check Box11",
    "Qty1", "SBMBDVBE Certification  if applicable",
}

# Strong "this is a CCHCS LPA IT Goods RFQ" signal — AcroForm field markers
# unique to the LPA template (see `cchcs_it_rfq_reytech_standard.yaml` +
# `fill_cchcs_it_rfq` dispatcher). Mutually exclusive with CCHCS_PACKET
# because the LPA is a different form with different field naming.
CCHCS_IT_RFQ_FIELD_NAMES = {
    "Supplier Address 1", "Supplier Address 2", "Supplier Email",
    "Extension TotalSubtotal", "Extension TotalSales Tax",
    "AMS 708 GenAI No", "Item Description1",
}

# Body/subject keywords that flag an LPA IT RFQ (classifier Pass 1, per
# feedback_ghost_record_heuristics "read email 3x" rule). At least one
# pair of LPA-specific + agency-specific tokens should be present.
CCHCS_IT_RFQ_BODY_PATTERNS = [
    r"\bLPA\s*#",
    r"LPA\s*IT\s*(?:Goods|Services)",
    r"Request\s+For\s+Quotation.{0,40}IT\s*(?:Goods|Services)",
    r"IT\s*Goods\s*and\s*Services",
]

# Strong "this is an AMS 704 worksheet" signals (text content)
AMS_704_HEADLINE_PATTERNS = [
    r"PRICE\s+CHECK\s+WORKSHEET",
    r"AMS\s*704",
    r"\bPrice\s+Check\s*#",
]

# Strong "this is a DocuSign PDF" signal
DOCUSIGN_PRODUCER = "docusign"
DOCUSIGN_ENVELOPE_PATTERN = r"Docusign\s*Envelope\s*ID"


def _has_lpa_body_signal(*text_blobs: str) -> bool:
    """Return True if any blob contains an LPA IT RFQ keyword pattern."""
    corpus = " ".join(t for t in text_blobs if t)
    if not corpus:
        return False
    for pat in CCHCS_IT_RFQ_BODY_PATTERNS:
        if re.search(pat, corpus, re.IGNORECASE):
            return True
    return False


# ── Main entry point ─────────────────────────────────────────────────────

def classify_request(
    attachments: List[str] = None,
    email_body: str = "",
    email_subject: str = "",
    email_sender: str = "",
) -> RequestClassification:
    """Classify a buyer request.

    Args:
        attachments: list of local file paths to the attachments
        email_body: plain-text body of the originating email (optional)
        email_subject: subject line (optional)
        email_sender: from-address of the email (optional)

    Returns: RequestClassification — never None, always at minimum
    `shape=SHAPE_UNKNOWN, agency="other", confidence=0.0`.
    """
    result = RequestClassification()
    result.reasons = []
    attachments = attachments or []

    # ── Shape detection — iterate through attachments ──
    primary_shape = SHAPE_EMAIL_ONLY
    primary_file = ""
    primary_file_type = ""
    producer = ""
    inline_text_samples: List[str] = []
    inline_form_field_sample: List[str] = []
    inline_field_count = 0
    page_count = 0
    docusign_detected = False

    for path in attachments:
        if not path or not os.path.exists(path):
            continue
        fname = os.path.basename(path)
        ext = os.path.splitext(fname)[1].lower()

        if ext == ".pdf":
            shape, info = _classify_pdf(path)
            if shape and shape != SHAPE_UNKNOWN:
                # Prefer packet > 704 > generic. CCHCS packet wins outright.
                if shape == SHAPE_CCHCS_PACKET:
                    primary_shape = shape
                    primary_file = fname
                    primary_file_type = "pdf"
                    producer = info.get("producer", "")
                    inline_field_count = info.get("field_count", 0)
                    page_count = info.get("page_count", 0)
                    inline_text_samples.append(info.get("text_sample", ""))
                    result.reasons.append(
                        f"pdf '{fname}': CCHCS packet (18+ pgs, 183 form fields, CCHCS headers)"
                    )
                    break  # packet is authoritative — skip the rest
                # Otherwise, keep looking but hold as candidate
                if primary_shape in (SHAPE_EMAIL_ONLY, SHAPE_UNKNOWN) or _shape_rank(shape) > _shape_rank(primary_shape):
                    primary_shape = shape
                    primary_file = fname
                    primary_file_type = "pdf"
                    producer = info.get("producer", "")
                    inline_field_count = info.get("field_count", 0)
                    page_count = info.get("page_count", 0)
                    inline_text_samples.append(info.get("text_sample", ""))
                    if shape == SHAPE_PC_704_PDF_DOCUSIGN:
                        docusign_detected = True
                    result.reasons.append(
                        f"pdf '{fname}': {shape} (producer={producer[:30]!r}, "
                        f"{inline_field_count} fields, {page_count} pgs)"
                    )
        elif ext == ".docx":
            shape, info = _classify_docx(path)
            if shape and shape != SHAPE_UNKNOWN:
                if primary_shape in (SHAPE_EMAIL_ONLY, SHAPE_UNKNOWN) or _shape_rank(shape) > _shape_rank(primary_shape):
                    primary_shape = shape
                    primary_file = fname
                    primary_file_type = "docx"
                    inline_text_samples.append(info.get("text_sample", ""))
                    result.reasons.append(
                        f"docx '{fname}': {shape} ({info.get('headline','')})"
                    )
        elif ext in (".xlsx", ".xls"):
            shape, info = _classify_xlsx(path)
            if shape and shape != SHAPE_UNKNOWN:
                if primary_shape in (SHAPE_EMAIL_ONLY, SHAPE_UNKNOWN) or _shape_rank(shape) > _shape_rank(primary_shape):
                    primary_shape = shape
                    primary_file = fname
                    primary_file_type = "xlsx"
                    inline_text_samples.append(info.get("text_sample", ""))
                    result.reasons.append(
                        f"xlsx '{fname}': {shape} ({info.get('headline','')})"
                    )
        elif ext in (".doc",):
            # Legacy .doc — treat as DOCX-ish, flag for conversion
            primary_shape = SHAPE_PC_704_DOCX
            primary_file = fname
            primary_file_type = "doc"
            result.reasons.append(
                f"doc '{fname}': legacy Word format, needs conversion"
            )

    result.shape = primary_shape
    result.primary_file = primary_file
    result.primary_file_type = primary_file_type
    result.producer_signature = producer
    result.needs_overlay_fill = docusign_detected or (
        primary_shape == SHAPE_PC_704_PDF_DOCUSIGN
    )
    result.meta["inline_field_count"] = inline_field_count
    result.meta["page_count"] = page_count

    # ── Agency detection — scan every available text source ──
    corpus_parts = [
        email_subject or "",
        email_body or "",
        email_sender or "",
    ]
    corpus_parts.extend(inline_text_samples)
    # Also scan attachment filenames for agency hints (VHC-WLA etc.)
    for path in attachments:
        corpus_parts.append(os.path.basename(path))
    corpus = " ".join(corpus_parts).upper()

    agency_matches = []
    for pattern, agency_key in AGENCY_KEYWORDS:
        if re.search(pattern, corpus, re.IGNORECASE):
            agency_matches.append(agency_key)

    # Institution prefix scan (handles 2-4 letter codes like CIW, VHC-WLA)
    for prefix, agency_key in INSTITUTION_TO_AGENCY.items():
        if re.search(rf"\b{re.escape(prefix)}\b", corpus):
            agency_matches.append(agency_key)

    # Shape-implies-agency fallback: a CCHCS packet / LPA IT RFQ by definition
    # IS a CCHCS request (no other agency uses these form layouts). Handles
    # the case where the page-1 logo is an image and no text-based agency
    # keyword matches.
    if result.shape == SHAPE_CCHCS_PACKET and "cchcs" not in agency_matches:
        agency_matches.append("cchcs")
        result.reasons.append("agency: cchcs (implied by cchcs_packet shape)")
    if result.shape == SHAPE_CCHCS_IT_RFQ and "cchcs" not in agency_matches:
        agency_matches.append("cchcs")
        result.reasons.append("agency: cchcs (implied by cchcs_it_rfq shape)")

    # Pick most-specific match: Barstow > CalVet > CCHCS > DSH > DGS
    if agency_matches:
        result.agency = _pick_most_specific_agency(agency_matches)
        if not any("implied" in r for r in result.reasons):
            result.reasons.append(
                f"agency: {result.agency} (matches: {', '.join(sorted(set(agency_matches)))})"
            )
    else:
        result.agency = "other"

    # ── Resolve required_forms from agency_config ──
    try:
        from src.core.agency_config import DEFAULT_AGENCY_CONFIGS
        cfg = DEFAULT_AGENCY_CONFIGS.get(result.agency, {})
        result.agency_name = cfg.get("name", "")
        raw_required = list(cfg.get("required_forms", []))
        # Narrow by classifier shape: an LPA / generic / email-only RFQ from a
        # packet-agency must not require 703B/704B/bidpkg that the buyer never
        # sent. See BUYER_TEMPLATE_FORMS + _SHAPE_BUYER_TEMPLATE_REQUIREMENTS.
        result.required_forms = filter_required_forms_by_shape(
            raw_required, result.shape
        )
        result.optional_forms = list(cfg.get("optional_forms", []))
        if len(result.required_forms) != len(raw_required):
            _dropped = set(raw_required) - set(result.required_forms)
            result.reasons.append(
                f"required_forms: dropped {sorted(_dropped)} — not applicable to shape={result.shape}"
            )
    except Exception as e:
        log.debug("agency_config lookup failed: %s", e)

    # ── Is this a "quote only" (PC) or "full package" (RFQ)? ──
    result.is_quote_only = result.shape in (
        SHAPE_PC_704_DOCX,
        SHAPE_PC_704_PDF_DOCUSIGN,
        SHAPE_PC_704_PDF_FILLABLE,
    )

    # ── Solicitation number extraction ──
    result.solicitation_number = _extract_solicitation(corpus, inline_text_samples)
    if result.solicitation_number:
        result.reasons.append(f"solicitation: {result.solicitation_number}")

    # ── Institution extraction (for agency confirmation / linking) ──
    result.institution = _extract_institution(corpus)
    if result.institution:
        result.reasons.append(f"institution: {result.institution}")

    # ── Confidence scoring ──
    result.confidence = _score_confidence(result, attachments, agency_matches)

    log.info(
        "classify_request: shape=%s agency=%s confidence=%.2f file=%s",
        result.shape, result.agency, result.confidence, primary_file or "(email only)"
    )
    return result


# ── Per-file-type classifiers ────────────────────────────────────────────

def _classify_pdf(path: str) -> Tuple[str, Dict[str, Any]]:
    """Inspect a PDF's form fields, metadata, and headline text to
    decide what shape it is. Returns (shape, info) where info has
    producer/field_count/page_count/text_sample for downstream use."""
    info: Dict[str, Any] = {
        "producer": "", "field_count": 0, "page_count": 0, "text_sample": "",
    }
    try:
        from pypdf import PdfReader
    except ImportError:
        return SHAPE_UNKNOWN, info

    try:
        reader = PdfReader(path)
        info["page_count"] = len(reader.pages)
        meta = reader.metadata or {}
        info["producer"] = str(meta.get("/Producer", "") or "")
        fields = reader.get_fields() or {}
        info["field_count"] = len(fields)
    except Exception as e:
        log.debug("pypdf read %s failed: %s", path, e)
        return SHAPE_UNKNOWN, info

    # Capture form field VALUES as additional text (Institution Name,
    # Requestor, etc. carry agency/institution signals even when the
    # page-1 CCHCS branding is an image)
    field_value_text = []
    for name, spec in fields.items():
        try:
            v = spec.get("/V", "") if isinstance(spec, dict) else ""
            if v and isinstance(v, str):
                field_value_text.append(v)
        except Exception:
            continue
    info["field_value_text"] = " ".join(field_value_text)

    field_names = set(fields.keys())

    # ── CCHCS packet detection (checked first — packet is a superset) ──
    # The 18-page CCHCS packet fixture contains ALL LPA field names AS WELL
    # AS packet-specific ones (Institution Name, Check Box11, larger page
    # count). Ordering: packet first so LPA doesn't steal packet-class
    # templates. LPA falls through to its own narrower check below.
    cchcs_field_overlap = len(field_names & CCHCS_PACKET_FIELD_NAMES)
    if (info["page_count"] >= 10 and cchcs_field_overlap >= 4) or info["field_count"] > 100:
        # Double-check with a headline scan so we don't mis-identify
        # some other 100+ field PDF
        text_sample = _pdf_headline(path)
        info["text_sample"] = text_sample + " " + info["field_value_text"]
        if ("Request For Quotation" in text_sample
                or "Non-Cloud" in text_sample
                or cchcs_field_overlap >= 5):
            return SHAPE_CCHCS_PACKET, info

    # ── CCHCS LPA IT Goods & Services RFQ detection ──
    # Standalone LPA templates are ~13 pages and ~100-180 fields — page
    # count is the primary discriminator from the full 18-page packet which
    # contains the LPA pages plus 703B/704B/bid-package pages. Real LPA
    # observed 2026-04-22 on RFQ 10840486: 13 pages, 181 fields.
    # Packet check above has already consumed 18+ page fixtures with
    # packet-specific markers. Anything reaching here is LPA-candidate.
    lpa_field_overlap = len(field_names & CCHCS_IT_RFQ_FIELD_NAMES)
    if (info["page_count"] < 16 and lpa_field_overlap >= 2):
        _lpa_text = _pdf_headline(path)
        info["text_sample"] = _lpa_text + " " + info["field_value_text"]
        # Belt-and-suspenders: body keyword confirms, OR strong fingerprint
        # (≥4 unique markers) is enough on its own.
        if _has_lpa_body_signal(_lpa_text, info["field_value_text"]) or lpa_field_overlap >= 4:
            return SHAPE_CCHCS_IT_RFQ, info

    # ── DocuSign PDF detection ──
    text_sample = _pdf_headline(path)
    info["text_sample"] = text_sample + " " + info["field_value_text"]
    is_docusign = (
        DOCUSIGN_PRODUCER in info["producer"].lower()
        or re.search(DOCUSIGN_ENVELOPE_PATTERN, text_sample, re.IGNORECASE)
    )
    is_704_by_text = any(
        re.search(p, text_sample, re.IGNORECASE) for p in AMS_704_HEADLINE_PATTERNS
    )

    if is_docusign and is_704_by_text:
        return SHAPE_PC_704_PDF_DOCUSIGN, info

    # ── Fillable 704 detection ──
    if is_704_by_text and info["field_count"] > 0:
        # Look for Row fields as confirmation
        row_fields = [n for n in field_names if re.match(r"(Row|QTYRow|UOMRow)\d+", n)]
        if len(row_fields) >= 3:
            return SHAPE_PC_704_PDF_FILLABLE, info

    # ── Fallback: non-704 PDF RFQ ──
    if info["page_count"] > 0:
        return SHAPE_GENERIC_RFQ_PDF, info

    return SHAPE_UNKNOWN, info


def _classify_docx(path: str) -> Tuple[str, Dict[str, Any]]:
    """Classify a DOCX by inspecting paragraphs + tables."""
    info: Dict[str, Any] = {"headline": "", "text_sample": ""}
    try:
        from docx import Document
    except ImportError:
        return SHAPE_UNKNOWN, info

    try:
        d = Document(path)
        # Headline is the first few non-empty paragraphs joined
        paras = [p.text.strip() for p in d.paragraphs if p.text.strip()]
        headline = " | ".join(paras[:5])
        info["headline"] = headline[:200]
        info["text_sample"] = headline

        # 704 detection via headline
        if any(re.search(p, headline, re.IGNORECASE) for p in AMS_704_HEADLINE_PATTERNS):
            return SHAPE_PC_704_DOCX, info

        # Table-based 704 detection (if header is in a table, not a paragraph)
        for t in d.tables[:3]:
            for row in t.rows[:3]:
                for cell in row.cells:
                    ct = cell.text.strip()
                    if any(re.search(p, ct, re.IGNORECASE) for p in AMS_704_HEADLINE_PATTERNS):
                        info["text_sample"] = info["text_sample"] + " | " + ct[:60]
                        return SHAPE_PC_704_DOCX, info

        # Fallback: generic DOCX RFQ
        return SHAPE_GENERIC_RFQ_DOCX, info
    except Exception as e:
        log.debug("docx read %s failed: %s", path, e)
        return SHAPE_UNKNOWN, info


def _classify_xlsx(path: str) -> Tuple[str, Dict[str, Any]]:
    """Classify a spreadsheet. Currently always SHAPE_GENERIC_RFQ_XLSX
    since there's no standardized spreadsheet 704 format. The text
    sample includes the first few non-empty cells for agency detection."""
    info: Dict[str, Any] = {"headline": "", "text_sample": ""}
    try:
        from openpyxl import load_workbook
    except ImportError:
        return SHAPE_UNKNOWN, info

    try:
        wb = load_workbook(path, read_only=True, data_only=True)
        text_chunks = []
        for sheet_name in wb.sheetnames[:3]:
            ws = wb[sheet_name]
            for row in ws.iter_rows(min_row=1, max_row=10, values_only=True):
                for cell in row:
                    if cell and isinstance(cell, str) and len(cell) > 2:
                        text_chunks.append(str(cell))
        headline = " | ".join(text_chunks[:20])
        info["headline"] = headline[:200]
        info["text_sample"] = headline
        return SHAPE_GENERIC_RFQ_XLSX, info
    except Exception as e:
        log.debug("xlsx read %s failed: %s", path, e)
        return SHAPE_UNKNOWN, info


# ── Helpers ──────────────────────────────────────────────────────────────

def _pdf_headline(path: str, max_chars: int = 500) -> str:
    """Extract the first few words of a PDF for text-based classification."""
    try:
        import pdfplumber
        with pdfplumber.open(path) as pdf:
            if not pdf.pages:
                return ""
            pg1 = pdf.pages[0]
            words = pg1.extract_words(keep_blank_chars=False)[:60]
            return " ".join(w.get("text", "") for w in words)[:max_chars]
    except Exception:
        return ""


_SHAPE_RANK = {
    SHAPE_CCHCS_PACKET: 100,
    SHAPE_CCHCS_IT_RFQ: 95,          # just below packet — both are strong CCHCS signals
    SHAPE_PC_704_PDF_FILLABLE: 80,
    SHAPE_PC_704_DOCX: 70,
    SHAPE_PC_704_PDF_DOCUSIGN: 60,
    SHAPE_GENERIC_RFQ_XLSX: 30,
    SHAPE_GENERIC_RFQ_PDF: 20,
    SHAPE_GENERIC_RFQ_DOCX: 20,
    SHAPE_EMAIL_ONLY: 10,
    SHAPE_UNKNOWN: 0,
}


def _shape_rank(shape: str) -> int:
    return _SHAPE_RANK.get(shape, 0)


_AGENCY_SPECIFICITY = {
    "calvet_barstow": 100,
    "calvet": 90,
    "cchcs": 80,
    "dsh": 70,
    "calfire": 60,
    "dgs": 50,
    "other": 0,
}


def _pick_most_specific_agency(matches: List[str]) -> str:
    """When multiple agency patterns match, pick the most specific one."""
    if not matches:
        return "other"
    return max(matches, key=lambda a: _AGENCY_SPECIFICITY.get(a, 0))


def _extract_solicitation(corpus: str, text_samples: List[str]) -> str:
    """Pull a solicitation number out of the corpus. Patterns:
    PREQ1234567, 10843276, RFQ-2026-XXXX, SOL-XXXX, etc."""
    all_text = corpus + " " + " ".join(text_samples)
    # Most specific first: PREQ + 6-8 digits
    m = re.search(r"\bPREQ(\d{6,10})\b", all_text, re.IGNORECASE)
    if m:
        return m.group(1)
    # 8-digit CCHCS solicitation numbers (e.g. 10843276)
    m = re.search(r"\b(\d{8})\b", all_text)
    if m:
        return m.group(1)
    # RFQ prefix
    m = re.search(r"\bRFQ[- ]([A-Z0-9-]{4,15})\b", all_text, re.IGNORECASE)
    if m:
        return m.group(1)
    return ""


def _extract_institution(corpus: str) -> str:
    """Best-effort extraction of the buyer institution name."""
    # Look for "Institution:" or "Delivery to:" patterns
    m = re.search(r"Institution\s*(?:Name|or HQ Program)[:\s]+([A-Z][A-Za-z0-9 &,\-]{3,60})",
                  corpus, re.IGNORECASE)
    if m:
        return m.group(1).strip().rstrip(",")
    # CCHCS pattern: "CA State Prison <Name>"
    m = re.search(r"(CA\s+State\s+Prison[- ][A-Za-z ]+)", corpus, re.IGNORECASE)
    if m:
        return m.group(1).strip()
    # VHC variants
    m = re.search(r"(VHC[- ](?:WLA|YV|CH|LB|FR|RED|VRV))", corpus, re.IGNORECASE)
    if m:
        return m.group(1).upper()
    return ""


def _score_confidence(
    result: RequestClassification,
    attachments: List[str],
    agency_matches: List[str],
) -> float:
    """Heuristic confidence score 0.0-1.0. Rewards strong signals
    (known shape + known agency + solicitation number) and penalizes
    fallbacks (unknown shape, other agency, no sol#)."""
    score = 0.0
    # Shape confidence
    if result.shape == SHAPE_CCHCS_PACKET:
        score += 0.40
    elif result.shape == SHAPE_CCHCS_IT_RFQ:
        score += 0.40  # LPA is as specific as packet — unique field namespace
    elif result.shape in (SHAPE_PC_704_DOCX, SHAPE_PC_704_PDF_FILLABLE):
        score += 0.35
    elif result.shape == SHAPE_PC_704_PDF_DOCUSIGN:
        score += 0.30
    elif result.shape in (SHAPE_GENERIC_RFQ_XLSX, SHAPE_GENERIC_RFQ_PDF, SHAPE_GENERIC_RFQ_DOCX):
        score += 0.20
    elif result.shape == SHAPE_EMAIL_ONLY:
        score += 0.10

    # Agency confidence
    if result.agency != "other":
        score += 0.25
        if len(set(agency_matches)) >= 2:
            score += 0.10  # corroborating signals

    # Solicitation number adds certainty
    if result.solicitation_number:
        score += 0.15

    # Institution adds a little
    if result.institution:
        score += 0.05

    # Attachments present
    if attachments:
        score += 0.05

    # Floor: a CCHCS packet with a confirmed agency match is a known
    # high-signal request even when the buyer forgot to include a
    # solicitation number on the cover page. Without this floor the
    # same packet dips to 0.70 and falls into manual review, which
    # is noisy for the #1 buyer we see.
    if result.shape == SHAPE_CCHCS_PACKET and result.agency != "other":
        score = max(score, 0.80)

    return min(1.0, round(score, 2))


def classify_enabled() -> bool:
    """Feature flag gate. When False, callers fall back to the legacy
    parallel PC/RFQ paths. Defaults to False during Phase 1 shipping —
    the classifier runs in shadow mode (logged but not dispatched).
    """
    try:
        from src.core.flags import get_flag
        return bool(get_flag("ingest.classifier_v2_enabled", False))
    except Exception:
        return False


__all__ = [
    "classify_request",
    "classify_enabled",
    "RequestClassification",
    # Shape constants
    "SHAPE_CCHCS_PACKET",
    "SHAPE_PC_704_DOCX",
    "SHAPE_PC_704_PDF_DOCUSIGN",
    "SHAPE_PC_704_PDF_FILLABLE",
    "SHAPE_GENERIC_RFQ_XLSX",
    "SHAPE_GENERIC_RFQ_PDF",
    "SHAPE_GENERIC_RFQ_DOCX",
    "SHAPE_EMAIL_ONLY",
    "SHAPE_UNKNOWN",
]
