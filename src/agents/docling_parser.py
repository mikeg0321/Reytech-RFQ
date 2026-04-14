"""
docling_parser.py — Intelligent document parser for procurement documents.

Parses PDFs, DOCX, and other documents into structured Markdown with table extraction.
Primary: docling library (97.9% table accuracy). Fallback: pypdf + python-docx.

V1: Parse and store. V2: Auto-detect doc type, competitor quote extraction, OCR.
"""
import json
import logging
import os
import time

log = logging.getLogger("reytech.docling_parser")

# Graceful degradation — docling is optional
try:
    from docling.document_converter import DocumentConverter
    HAS_DOCLING = True
except ImportError:
    HAS_DOCLING = False
    log.info("docling not installed — using pypdf/python-docx fallback")

ALLOWED_EXTENSIONS = {".pdf", ".docx", ".doc", ".xlsx"}
MAX_FILE_SIZE = 20 * 1024 * 1024  # 20MB
PARSE_TIMEOUT = 30  # seconds


# ═══════════════════════════════════════════════════════════════════════════
# PARSING ENGINES
# ═══════════════════════════════════════════════════════════════════════════

def _parse_with_docling(file_path: str) -> dict:
    """Primary parser using IBM docling — high-accuracy table extraction."""
    converter = DocumentConverter()
    result = converter.convert(file_path)

    markdown = result.document.export_to_markdown()

    # Extract tables as structured data
    tables = []
    for table in result.document.tables:
        table_data = {
            "caption": getattr(table, "caption", ""),
            "rows": [],
        }
        if hasattr(table, "export_to_dataframe"):
            try:
                df = table.export_to_dataframe()
                table_data["headers"] = list(df.columns)
                table_data["rows"] = df.values.tolist()
            except Exception as _e:
                log.debug("suppressed: %s", _e)
        tables.append(table_data)

    page_count = getattr(result.document, "num_pages", 0)
    if not page_count and hasattr(result, "pages"):
        page_count = len(result.pages)

    return {
        "markdown": markdown,
        "tables": tables,
        "page_count": page_count,
        "engine": "docling",
    }


def _parse_with_pypdf(file_path: str) -> dict:
    """Fallback PDF parser using pypdf (already in requirements.txt)."""
    from pypdf import PdfReader

    reader = PdfReader(file_path)
    pages_text = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        pages_text.append(f"## Page {i + 1}\n\n{text}")

    markdown = "\n\n---\n\n".join(pages_text)
    return {
        "markdown": markdown,
        "tables": [],  # pypdf doesn't extract tables structurally
        "page_count": len(reader.pages),
        "engine": "pypdf",
    }


def _parse_with_docx(file_path: str) -> dict:
    """Fallback DOCX parser using python-docx (already in requirements.txt)."""
    from docx import Document

    doc = Document(file_path)
    parts = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        # Detect heading styles
        if para.style and para.style.name and para.style.name.startswith("Heading"):
            level = para.style.name.replace("Heading ", "")
            try:
                hashes = "#" * int(level)
            except ValueError:
                hashes = "##"
            parts.append(f"{hashes} {text}")
        else:
            parts.append(text)

    # Extract tables
    tables = []
    for table in doc.tables:
        rows = []
        headers = []
        for i, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            if i == 0:
                headers = cells
            else:
                rows.append(cells)
        tables.append({"headers": headers, "rows": rows, "caption": ""})

    # Add table markdown
    for tbl in tables:
        if tbl["headers"]:
            parts.append("\n| " + " | ".join(tbl["headers"]) + " |")
            parts.append("| " + " | ".join(["---"] * len(tbl["headers"])) + " |")
            for row in tbl["rows"]:
                parts.append("| " + " | ".join(row) + " |")

    return {
        "markdown": "\n\n".join(parts),
        "tables": tables,
        "page_count": 0,  # DOCX doesn't have fixed pages
        "engine": "python-docx",
    }


def _parse_with_xlsx(file_path: str) -> dict:
    """Fallback XLSX parser using openpyxl (already in requirements.txt)."""
    from openpyxl import load_workbook

    wb = load_workbook(file_path, read_only=True, data_only=True)
    parts = []
    tables = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        parts.append(f"## Sheet: {sheet_name}\n")
        rows = []
        headers = []
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            cells = [str(c) if c is not None else "" for c in row]
            if i == 0:
                headers = cells
            else:
                rows.append(cells)

        if headers:
            tables.append({"headers": headers, "rows": rows, "caption": sheet_name})
            parts.append("| " + " | ".join(headers) + " |")
            parts.append("| " + " | ".join(["---"] * len(headers)) + " |")
            for row in rows[:100]:  # Cap at 100 rows for markdown
                parts.append("| " + " | ".join(row) + " |")

    wb.close()
    return {
        "markdown": "\n\n".join(parts),
        "tables": tables,
        "page_count": len(wb.sheetnames),
        "engine": "openpyxl",
    }


# ═══════════════════════════════════════════════════════════════════════════
# PUBLIC API
# ═══════════════════════════════════════════════════════════════════════════

def validate_file(file_path: str) -> tuple:
    """Validate file for parsing. Returns (ok: bool, error: str)."""
    if not os.path.isfile(file_path):
        return False, "File not found"

    ext = os.path.splitext(file_path)[1].lower()
    if ext not in ALLOWED_EXTENSIONS:
        return False, f"Unsupported file type: {ext}. Allowed: {', '.join(ALLOWED_EXTENSIONS)}"

    size = os.path.getsize(file_path)
    if size > MAX_FILE_SIZE:
        return False, f"File too large: {size / 1024 / 1024:.1f}MB (max {MAX_FILE_SIZE / 1024 / 1024:.0f}MB)"

    if size == 0:
        return False, "File is empty"

    return True, ""


def parse_document(file_path: str) -> dict:
    """Parse a document into structured Markdown + tables.

    Returns: {"ok", "markdown", "tables", "metadata", "page_count", "duration_ms", "engine"}
    """
    ok, err = validate_file(file_path)
    if not ok:
        return {"ok": False, "error": err}

    ext = os.path.splitext(file_path)[1].lower()
    start = time.time()

    try:
        if ext == ".pdf":
            if HAS_DOCLING:
                try:
                    result = _parse_with_docling(file_path)
                except Exception as e:
                    log.warning("Docling parse failed, falling back to pypdf: %s", e)
                    result = _parse_with_pypdf(file_path)
            else:
                result = _parse_with_pypdf(file_path)

        elif ext in (".docx", ".doc"):
            if HAS_DOCLING:
                try:
                    result = _parse_with_docling(file_path)
                except Exception as e:
                    log.warning("Docling DOCX parse failed, falling back: %s", e)
                    result = _parse_with_docx(file_path)
            else:
                result = _parse_with_docx(file_path)

        elif ext == ".xlsx":
            result = _parse_with_xlsx(file_path)

        else:
            return {"ok": False, "error": f"No parser for {ext}"}

        duration_ms = int((time.time() - start) * 1000)

        log.info("Parsed %s: %d pages, %d tables, %dms (%s)",
                 os.path.basename(file_path), result.get("page_count", 0),
                 len(result.get("tables", [])), duration_ms, result.get("engine", "unknown"))

        return {
            "ok": True,
            "markdown": result["markdown"],
            "tables": result.get("tables", []),
            "metadata": {
                "filename": os.path.basename(file_path),
                "extension": ext,
                "size_bytes": os.path.getsize(file_path),
                "engine": result.get("engine", "unknown"),
            },
            "page_count": result.get("page_count", 0),
            "duration_ms": duration_ms,
        }

    except Exception as e:
        log.error("Document parse failed for %s: %s", file_path, e, exc_info=True)
        return {"ok": False, "error": str(e)}


# ═══════════════════════════════════════════════════════════════════════════
# DATABASE OPERATIONS
# ═══════════════════════════════════════════════════════════════════════════

def save_parsed_document(filename: str, markdown: str, tables: list,
                         metadata: dict, linked_rfq_id: str = None,
                         linked_pc_id: str = None, doc_type: str = "unknown",
                         page_count: int = 0, duration_ms: int = 0) -> int:
    """Save parsed document to DB. Returns doc ID or -1 on error."""
    try:
        from src.core.db import get_db
        with get_db() as conn:
            cursor = conn.execute(
                """INSERT INTO parsed_documents
                   (filename, doc_type, parsed_markdown, parsed_tables_json,
                    metadata_json, linked_rfq_id, linked_pc_id,
                    parse_duration_ms, page_count)
                   VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)""",
                (filename, doc_type, markdown, json.dumps(tables),
                 json.dumps(metadata), linked_rfq_id, linked_pc_id,
                 duration_ms, page_count)
            )
            doc_id = cursor.lastrowid
            log.info("Saved parsed document: id=%d, filename=%s, type=%s",
                     doc_id, filename, doc_type)
            return doc_id
    except Exception as e:
        log.error("Failed to save parsed document %s: %s", filename, e, exc_info=True)
        return -1


def get_parsed_document(doc_id: int) -> dict:
    """Load a parsed document from DB."""
    try:
        from src.core.db import get_db
        with get_db() as conn:
            row = conn.execute(
                "SELECT * FROM parsed_documents WHERE id = ?", (doc_id,)
            ).fetchone()
            if not row:
                return {"ok": False, "error": "Document not found"}
            d = dict(row)
            # Parse JSON fields
            for field in ("parsed_tables_json", "metadata_json"):
                if d.get(field):
                    try:
                        d[field] = json.loads(d[field])
                    except (json.JSONDecodeError, TypeError) as _e:
                        log.debug("suppressed: %s", _e)
            d["ok"] = True
            return d
    except Exception as e:
        log.error("Failed to load parsed document %d: %s", doc_id, e)
        return {"ok": False, "error": str(e)}


def list_parsed_documents(limit: int = 50, doc_type: str = None) -> list:
    """List parsed documents, newest first."""
    try:
        from src.core.db import get_db
        with get_db() as conn:
            if doc_type:
                rows = conn.execute(
                    "SELECT id, filename, doc_type, page_count, parse_duration_ms, uploaded_at, status "
                    "FROM parsed_documents WHERE doc_type = ? ORDER BY uploaded_at DESC LIMIT ?",
                    (doc_type, limit)
                ).fetchall()
            else:
                rows = conn.execute(
                    "SELECT id, filename, doc_type, page_count, parse_duration_ms, uploaded_at, status "
                    "FROM parsed_documents ORDER BY uploaded_at DESC LIMIT ?",
                    (limit,)
                ).fetchall()
            return [dict(r) for r in rows]
    except Exception as e:
        log.error("Failed to list parsed documents: %s", e)
        return []
