"""
Office document text extractor for procurement parsing pipeline.

Extracts text/tables from XLS, XLSX, DOC, DOCX files so they can be
fed into the same Claude API extraction used by the vision parser.

Supported formats:
- XLSX: openpyxl (already in requirements)
- XLS:  xlrd (optional — falls back to error message)
- DOCX: python-docx (optional — falls back to error message)
- DOC:  not natively supported — user should convert to DOCX
"""

import os
import logging

log = logging.getLogger("reytech.doc_converter")

# Accepted office extensions (lowercase, with dot)
OFFICE_EXTS = (".xls", ".xlsx", ".doc", ".docx")

# All parseable extensions (office + pdf + image)
ALL_UPLOAD_EXTS = (
    ".pdf", ".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp", ".tiff", ".tif",
    ".xls", ".xlsx", ".doc", ".docx",
)

# HTML accept string for file inputs
ACCEPT_ALL = (
    ".pdf,.png,.jpg,.jpeg,.gif,.webp,.bmp,.tiff,"
    ".xls,.xlsx,.doc,.docx"
)


def is_office_doc(filename: str) -> bool:
    """Check if filename has an office document extension."""
    return os.path.splitext(filename)[1].lower() in OFFICE_EXTS


def extract_text(file_path: str) -> str:
    """Extract text content from an office document.

    Returns a plain-text representation suitable for Claude API extraction.
    Raises ValueError if the format is unsupported or a required library is missing.
    """
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".xlsx":
        return _extract_xlsx(file_path)
    elif ext == ".xls":
        return _extract_xls(file_path)
    elif ext == ".docx":
        return _extract_docx(file_path)
    elif ext == ".doc":
        raise ValueError(
            "Legacy .doc format is not supported. Please save as .docx and re-upload."
        )
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def _extract_xlsx(file_path: str) -> str:
    """Extract text from XLSX using openpyxl."""
    try:
        from openpyxl import load_workbook
    except ImportError:
        raise ValueError("openpyxl is not installed — cannot read XLSX files")

    wb = load_workbook(file_path, read_only=True, data_only=True)
    parts = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            continue

        parts.append(f"=== Sheet: {sheet_name} ===")

        for row in rows:
            # Skip fully empty rows
            vals = [str(c) if c is not None else "" for c in row]
            if not any(v.strip() for v in vals):
                continue
            parts.append("\t".join(vals))

    wb.close()

    if not parts:
        raise ValueError("XLSX file has no data in any sheet")

    text = "\n".join(parts)
    log.info("Extracted %d chars from XLSX (%d sheets)", len(text), len(wb.sheetnames))
    return text


def _extract_xls(file_path: str) -> str:
    """Extract text from legacy XLS using xlrd."""
    try:
        import xlrd
    except ImportError:
        raise ValueError(
            "xlrd is not installed — cannot read legacy .xls files. "
            "Please save as .xlsx and re-upload."
        )

    wb = xlrd.open_workbook(file_path)
    parts = []

    for sheet_idx in range(wb.nsheets):
        ws = wb.sheet_by_index(sheet_idx)
        if ws.nrows == 0:
            continue

        parts.append(f"=== Sheet: {ws.name} ===")

        for row_idx in range(ws.nrows):
            vals = []
            for col_idx in range(ws.ncols):
                cell = ws.cell(row_idx, col_idx)
                vals.append(str(cell.value) if cell.value is not None else "")
            if not any(v.strip() for v in vals):
                continue
            parts.append("\t".join(vals))

    if not parts:
        raise ValueError("XLS file has no data in any sheet")

    text = "\n".join(parts)
    log.info("Extracted %d chars from XLS (%d sheets)", len(text), wb.nsheets)
    return text


def _extract_docx(file_path: str) -> str:
    """Extract text from DOCX using python-docx.

    Detects AMS 704 Price Check DOCX format and produces structured output
    that clearly separates header metadata from line items, preventing the
    AI parser from treating header fields as product items.
    """
    try:
        import docx
    except ImportError:
        raise ValueError(
            "python-docx is not installed — cannot read DOCX files. "
            "Please convert to PDF and re-upload."
        )

    doc = docx.Document(file_path)

    # Detect AMS 704 format from paragraphs
    para_texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    is_704 = any("PRICE CHECK" in t.upper() or "AMS 704" in t.upper()
                 for t in para_texts)

    if is_704 and doc.tables:
        return _extract_docx_704(doc, para_texts)

    # Generic extraction for non-704 DOCX files
    parts = []
    for text in para_texts:
        parts.append(text)
    for table in doc.tables:
        parts.append("")
        for row in table.rows:
            vals = [cell.text.strip() for cell in row.cells]
            parts.append("\t".join(vals))

    if not parts:
        raise ValueError("DOCX file has no text or table content")

    text = "\n".join(parts)
    log.info("Extracted %d chars from DOCX (%d paragraphs, %d tables)",
             len(text), len(doc.paragraphs), len(doc.tables))
    return text


def _extract_docx_704(doc, para_texts: list) -> str:
    """Extract structured text from an AMS 704 Price Check DOCX.

    Identifies header metadata table, items tables, and footer/totals table,
    producing clearly labeled sections so the AI parser only extracts real items.
    """
    header_info = {}
    items_tables = []
    skipped_tables = 0

    for table in doc.tables:
        first_row_text = " ".join(
            cell.text.strip() for cell in table.rows[0].cells
        ).upper()

        # Items table: header row contains ITEM # and QTY
        if "ITEM #" in first_row_text and "QTY" in first_row_text:
            items_tables.append(table)

        # Header metadata table: contains PRICE CHECK # or Requestor
        elif "PRICE CHECK" in first_row_text or "DUE DATE" in first_row_text:
            # Extract header fields from the metadata table
            header_info.update(_parse_704_header_table(table))
            skipped_tables += 1

        # Footer/totals table: contains Subtotal or TOTAL PRICE
        elif "SUBTOTAL" in first_row_text or "TOTAL PRICE" in first_row_text:
            skipped_tables += 1

        # Supplier info / notes table — skip
        elif "SUPPLIER" in first_row_text and "NOTE" in first_row_text:
            skipped_tables += 1

        else:
            skipped_tables += 1

    # Build structured output
    parts = ["=== AMS 704 PRICE CHECK WORKSHEET ===", ""]

    # Header section
    if header_info:
        parts.append("=== HEADER (metadata — NOT line items) ===")
        for k, v in header_info.items():
            parts.append(f"{k}: {v}")
        parts.append("")

    # Items section — the only section the AI should extract items from
    if items_tables:
        parts.append("=== LINE ITEMS (extract items ONLY from this section) ===")
        for tbl_idx, table in enumerate(items_tables):
            merged_items = _merge_704_table_rows(table)
            for item in merged_items:
                parts.append(item)
            parts.append("")  # gap between pages
    else:
        log.warning("704 DOCX: no items tables found in %d tables", len(doc.tables))

    text = "\n".join(parts)
    log.info("704 DOCX extracted: header=%d fields, %d items tables, %d skipped tables, %d chars",
             len(header_info), len(items_tables), skipped_tables, len(text))
    return text


def _parse_704_header_table(table) -> dict:
    """Extract header fields from the AMS 704 header metadata table.

    The header table has merged cells with patterns like:
    'Requestor\\nCarolyn Montgomery' or 'Delivery Zip Code\\n92880'
    """
    header = {}
    field_map = {
        "requestor": "Requestor",
        "institution": "Institution",
        "delivery zip": "Delivery Zip Code",
        "phone": "Phone Number",
        "date of request": "Date of Request",
        "price check": "Price Check Number",
        "due date": "Due Date",
    }

    for row in table.rows:
        for cell in row.cells:
            cell_text = cell.text.strip()
            if not cell_text:
                continue
            cell_lower = cell_text.lower()

            # Skip pure label rows (SUPPLIER INFORMATION, COMPANY NAME, etc.)
            if cell_lower.startswith("s u p p l i e r"):
                continue
            if cell_lower in ("company name", "address", "signature and date",
                              "discount offered", "phone number", "e-mail address",
                              "company representative (print name)",
                              "delivery date and time (aro)",
                              "date price check expires",
                              "certified sb/mb #", "certified dvbe #"):
                continue

            # Look for "Label\nValue" pattern
            for key, label in field_map.items():
                if key in cell_lower:
                    lines = cell_text.split("\n")
                    if len(lines) >= 2:
                        value = lines[-1].strip()
                        if value and value.lower() != key:
                            header[label] = value
                    break

    return header


def _merge_704_table_rows(table) -> list:
    """Merge continuation rows in a 704 items table.

    AMS 704 DOCX items span 2 rows:
    Row N:   ITEM# | QTY | UOM | QTY_PER_UOM | Description | Substituted | Price | Extension
    Row N+1: (empty) | ... | ... | ... | Additional detail | Costco ref | ... | ...

    Returns list of formatted item strings.
    """
    rows = list(table.rows)
    if not rows:
        return []

    # First row is the header — include it for context
    header_cells = [cell.text.strip().replace("\n", " ") for cell in rows[0].cells]
    items = ["\t".join(header_cells)]

    i = 1
    while i < len(rows):
        cells = [cell.text.strip() for cell in rows[i].cells]
        item_num = cells[0] if cells else ""

        # Skip fully empty rows
        if not any(c.strip() for c in cells):
            i += 1
            continue

        # Skip footer rows (ENTER GRAND TOTAL...)
        joined = " ".join(cells).upper()
        if "GRAND TOTAL" in joined or "ENTER GRAND" in joined:
            i += 1
            continue

        if item_num and item_num.isdigit():
            # Primary item row — check if next row is a continuation
            desc = cells[4] if len(cells) > 4 else ""
            sub = cells[5] if len(cells) > 5 else ""

            # Look ahead for continuation row(s)
            j = i + 1
            while j < len(rows):
                next_cells = [cell.text.strip() for cell in rows[j].cells]
                next_item = next_cells[0] if next_cells else ""
                next_joined = " ".join(next_cells)

                # If next row has an item number or is empty or is footer, stop
                if next_item.isdigit():
                    break
                if not any(c.strip() for c in next_cells):
                    j += 1
                    continue
                if "GRAND TOTAL" in next_joined.upper():
                    j += 1
                    continue

                # Continuation row — append description and substituted item
                cont_desc = next_cells[4] if len(next_cells) > 4 else ""
                cont_sub = next_cells[5] if len(next_cells) > 5 else ""
                if cont_desc:
                    desc = desc + " | " + cont_desc if desc else cont_desc
                if cont_sub:
                    sub = sub + " | " + cont_sub if sub else cont_sub
                j += 1

            # Build merged item line
            merged = [
                item_num,
                cells[1] if len(cells) > 1 else "",   # QTY
                cells[2] if len(cells) > 2 else "",   # UOM
                cells[3] if len(cells) > 3 else "",   # QTY PER UOM
                desc,                                   # Merged description
                sub,                                    # Merged substituted item
                cells[6] if len(cells) > 6 else "",   # PRICE PER UNIT
                cells[7] if len(cells) > 7 else "",   # EXTENSION
            ]
            items.append("\t".join(merged))
            i = j
        else:
            # Orphan row without item number — skip (usually already merged)
            i += 1

    return items


# ═══════════════════════════════════════════════════════════════════════
# DOCX → PDF conversion via LibreOffice headless
# ═══════════════════════════════════════════════════════════════════════

import subprocess
import shutil


def convert_to_pdf(file_path: str, output_dir: str = None) -> str:
    """Convert a DOCX/DOC/XLS/XLSX file to PDF using LibreOffice headless.

    Returns the path to the generated PDF, or raises RuntimeError on failure.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Source file not found: {file_path}")

    if output_dir is None:
        output_dir = os.path.dirname(file_path)
    os.makedirs(output_dir, exist_ok=True)

    # Find LibreOffice binary
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        # Common paths
        for candidate in ["/usr/bin/libreoffice", "/usr/bin/soffice",
                          "C:/Program Files/LibreOffice/program/soffice.exe"]:
            if os.path.exists(candidate):
                soffice = candidate
                break
    if not soffice:
        raise RuntimeError("LibreOffice not found — cannot convert office docs to PDF")

    cmd = [
        soffice, "--headless", "--norestore", "--convert-to", "pdf",
        "--outdir", output_dir, file_path
    ]
    log.info("convert_to_pdf: %s", " ".join(cmd))
    try:
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
        if result.returncode != 0:
            raise RuntimeError(f"LibreOffice conversion failed: {result.stderr[:300]}")
    except subprocess.TimeoutExpired:
        raise RuntimeError("LibreOffice conversion timed out (60s)")

    # Find the output PDF
    base = os.path.splitext(os.path.basename(file_path))[0]
    pdf_path = os.path.join(output_dir, base + ".pdf")
    if not os.path.exists(pdf_path):
        raise RuntimeError(f"Conversion produced no PDF (expected: {pdf_path})")

    log.info("convert_to_pdf: %s → %s (%d bytes)",
             os.path.basename(file_path), os.path.basename(pdf_path),
             os.path.getsize(pdf_path))
    return pdf_path


def can_convert_to_pdf() -> bool:
    """Check if LibreOffice is available for DOCX→PDF conversion."""
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if soffice:
        return True
    for candidate in ["/usr/bin/libreoffice", "/usr/bin/soffice"]:
        if os.path.exists(candidate):
            return True
    return False


# ═══════════════════════════════════════════════════════════════════════
# Regex fallback: parse simple item lists from extracted text
# ═══════════════════════════════════════════════════════════════════════

import re


def _parse_704_structured_text(text: str) -> list:
    """Parse items from AMS 704 structured text produced by _extract_docx_704.

    Expects tab-separated lines in the LINE ITEMS section with columns:
    ITEM# | QTY | UOM | QTY_PER_UOM | DESCRIPTION | SUBSTITUTED | PRICE | EXTENSION
    """
    items = []
    in_items = False

    for line in text.split("\n"):
        line = line.strip()

        # Enter items section
        if line.startswith("=== LINE ITEMS"):
            in_items = True
            continue

        # Exit items section on next section header
        if in_items and line.startswith("==="):
            in_items = False
            continue

        if not in_items or not line:
            continue

        # Skip table header rows
        if "ITEM #" in line.upper() and "QTY" in line.upper():
            continue

        cols = line.split("\t")
        if len(cols) < 5:
            continue

        item_num = cols[0].strip()
        if not item_num or not item_num.isdigit():
            continue

        qty_str = cols[1].strip()
        try:
            qty = int(qty_str)
        except (ValueError, TypeError):
            qty = 1

        uom = cols[2].strip() or "each"
        qpu_raw = cols[3].strip()

        # Parse qty_per_uom — may be "12", "38oz.", "90oz", etc.
        qpu = 1
        qpu_match = re.match(r'(\d+)', qpu_raw)
        if qpu_match:
            qpu = int(qpu_match.group(1))
            if qpu < 1:
                qpu = 1

        desc = cols[4].strip() if len(cols) > 4 else ""
        sub = cols[5].strip() if len(cols) > 5 else ""

        # Use substituted item as primary description if it's more detailed
        # (common in 704 DOCX: description is short, substituted has full Costco name)
        final_desc = desc
        if sub and len(sub) > len(desc):
            final_desc = sub
        elif sub:
            final_desc = desc + " | " + sub

        # Extract Costco/supplier item number from both description columns
        part_number = ""
        for _col_text in (desc, sub, final_desc):
            item_id_match = re.search(r'Item\s*#?\s*(\d{4,7})', _col_text)
            if item_id_match:
                part_number = item_id_match.group(1)
                break

        items.append({
            "line_number": len(items) + 1,
            "item_number": item_num,
            "qty": qty,
            "uom": uom.lower(),
            "qty_per_uom": qpu,
            "description": final_desc,
            "part_number": part_number,
            "item_link": "",
            "row_index": len(items) + 1,  # 1-based for form field mapping (Row1, Row2, ...)
        })

    if items:
        log.info("704 structured text parsed %d items", len(items))
    return items


def parse_items_from_text(text: str) -> list:
    """Parse item descriptions + quantities from plain text (regex fallback).

    Handles common patterns like:
        Product description here
        Qty: 7
        ASIN = B0CZRF2DZR

        Another product description
        Qty = 4 Each

    Returns list of dicts in the standard line_items format.
    """
    if not text or len(text.strip()) < 10:
        return []

    # Detect AMS 704 structured text (from _extract_docx_704)
    if "=== LINE ITEMS" in text:
        return _parse_704_structured_text(text)

    lines = [l.strip() for l in text.split("\n")]
    items = []
    current_desc = ""
    current_qty = 1
    current_uom = "each"
    current_qpu = 1
    current_part = ""
    current_asin = ""

    # Patterns
    qty_pat = re.compile(r'^(?:qty|quantity)\s*[:=]\s*(\d+)\s*(.*)', re.IGNORECASE)
    qpu_pat = re.compile(r'^(?:qty\s*per\s*(?:uom|unit)|pack\s*(?:size|qty)|units?\s*per)\s*[:=]\s*(\d+)', re.IGNORECASE)
    asin_pat = re.compile(r'^ASIN\s*[:=]\s*(\S+)', re.IGNORECASE)
    upc_pat = re.compile(r'^UPC\s*[:=]\s*(\S+)', re.IGNORECASE)
    model_pat = re.compile(r'^(?:Model\s*(?:Number|#|No\.?)?|MFG\s*#?|Part\s*(?:Number|#|No\.?))\s*[:=]\s*(\S+)', re.IGNORECASE)
    unit_count_pat = re.compile(r'^Unit\s*Count\s*[:=]\s*(\d+)', re.IGNORECASE)
    # Skip lines that are metadata (product dimensions, color, style, etc.)
    meta_pat = re.compile(r'^(?:Product Dimensions|Color|Style|Base Material|Top Material|'
                          r'Finish Type|Special Feature|Brand|Item Weight|Manufacturer)\b', re.IGNORECASE)

    def _flush():
        nonlocal current_desc, current_qty, current_uom, current_qpu, current_part, current_asin
        if current_desc and len(current_desc) >= 5:
            # Detect pack size from description if not explicitly set
            qpu = current_qpu
            if qpu <= 1:
                _d = current_desc.upper()
                for _p in [r'PACK\s+(?:OF\s+)?(\d+)', r'(\d+)\s*[-/]?\s*(?:PACK|PK|COUNT|CT)',
                           r'(\d+)\s+PER\s+(?:BOX|CASE|PACK)', r'(\d+)\s*(?:PC|PCS)\b']:
                    _m = re.search(_p, _d)
                    if _m:
                        _v = int(_m.group(1))
                        if 1 < _v < 10000:
                            qpu = _v
                            break
            items.append({
                "line_number": len(items) + 1,
                "item_number": current_part or current_asin or str(len(items) + 1),
                "qty": current_qty,
                "uom": current_uom,
                "qty_per_uom": qpu,
                "description": current_desc,
                "part_number": current_part or current_asin,
                "item_link": "",
                "row_index": len(items) + 1,  # 1-based for form field mapping
            })
        current_desc = ""
        current_qty = 1
        current_uom = "each"
        current_qpu = 1
        current_part = ""
        current_asin = ""

    for line in lines:
        if not line:
            continue

        # Skip table metadata lines
        if meta_pat.match(line):
            continue

        # Check for qty_per_uom / unit count line
        m = qpu_pat.match(line)
        if m:
            current_qpu = max(1, int(m.group(1)))
            continue
        m = unit_count_pat.match(line)
        if m:
            current_qpu = max(1, int(m.group(1)))
            continue

        # Check for qty line
        m = qty_pat.match(line)
        if m:
            current_qty = int(m.group(1))
            uom_text = m.group(2).strip().lower()
            if uom_text and uom_text in ("each", "ea", "pack", "set", "box", "case", "pair"):
                current_uom = uom_text
            continue

        # Check for ASIN
        m = asin_pat.match(line)
        if m:
            current_asin = m.group(1)
            continue

        # Check for UPC / Model / Part number
        m = upc_pat.match(line) or model_pat.match(line)
        if m:
            current_part = m.group(1)
            continue

        # If line is long enough to be a description (>20 chars) and we already
        # have a description buffered, flush the previous item first
        if len(line) > 20 and current_desc:
            _flush()

        # Buffer as description (or append to short existing desc)
        if not current_desc:
            current_desc = line
        elif len(current_desc) < 20:
            current_desc += " " + line

    # Flush last item
    _flush()

    if items:
        log.info("Regex fallback parsed %d items from text", len(items))
    return items
